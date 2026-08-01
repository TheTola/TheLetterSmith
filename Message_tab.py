# ===============================
# File: Message_tab.py
# Purpose: Message authoring + render to message.png
# ===============================

from __future__ import annotations

import html as _html
import os
import re
import json
import shutil
import math
from pathlib import Path
from typing import Optional

# Optional converters (prefer mammoth for DOCX → HTML)
try:
    import mammoth  # type: ignore
except Exception:
    mammoth = None

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None  # type: ignore

try:
    from PyPDF2 import PdfReader  # type: ignore
except Exception:
    PdfReader = None  # type: ignore

from PySide6 import QtCore, QtGui, QtWidgets
from PySide6.QtCore import Qt, QSizeF, Signal
from PySide6.QtGui import (
    QTextDocument,
    QImage,
    QPainter,
    QPixmap,
    QFont,
)

from Editor import Editor
from image_button import ArtworkButton
from project_sync import file_fingerprint
from message_history import (
    delete_revision,
    list_revisions,
    restore_revision,
)
from message_format import normalize_ultralinks_in_document
from message_html import is_lettersmith_message_html
from project_paths import ProjectPathResolver
from project_save import ProjectNotReadyError, ProjectSaveService
from project_state import ProjectStateController
from settings_store import (
    SettingsStore,
    normalize_published_page_url,
)
from config import (
    SETTINGS_FILE,
    PUBLISHED_PAGE_URL_KEY,
    USER_PAGES_DIR,
    MESSAGE_HTML_FILE,
    MESSAGE_IMAGE_FILE,
)


IDENTITY_LOCK_KEYS = {
    "title": "recipient_title_locked",
    "recipient": "recipient_name_locked",
    "published_url": "published_page_url_locked",
}


class IdentityLineEdit(QtWidgets.QLineEdit):
    """Line edit that requires a double-click before editing a committed value."""

    double_clicked = Signal()

    def mouseDoubleClickEvent(self, event: QtGui.QMouseEvent) -> None:
        if self.isReadOnly():
            self.setReadOnly(False)
            self.double_clicked.emit()
        super().mouseDoubleClickEvent(event)

# ─────────────────────────────────────────────────────────────────────────────
# Utilities
# ─────────────────────────────────────────────────────────────────────────────

def _atomic_save_image(img: QImage, path: str) -> bool:
    """
    Windows-safe atomic PNG save using QtCore.QSaveFile.
    Uses QBuffer + QByteArray to encode first, then commits atomically.
    """
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)

    # Encode into memory
    ba = QtCore.QByteArray()
    buf = QtCore.QBuffer(ba)
    if not buf.open(QtCore.QIODevice.WriteOnly):
        return False

    ok = img.save(buf, "PNG")
    buf.close()

    if not ok or ba.isEmpty():
        return False

    # Commit with retries (Windows transient locks)
    for _ in range(8):
        sf = QtCore.QSaveFile(str(p))
        if not sf.open(QtCore.QIODevice.WriteOnly):
            QtCore.QThread.msleep(80)
            continue

        written = sf.write(ba)
        if written != ba.size():
            sf.cancelWriting()
            QtCore.QThread.msleep(80)
            continue

        if sf.commit():
            return True

        sf.cancelWriting()
        QtCore.QThread.msleep(80)

    return False


def _plain_text_from_html(raw_html: str) -> str:
    doc = QTextDocument()
    doc.setHtml(raw_html or "")
    return doc.toPlainText()


def _strip_html_for_word_count(html: str) -> str:
    return _plain_text_from_html(html)


def _word_count(text_like: str) -> int:
    return len(re.findall(r"\b[\w’'-]+\b", text_like or "", flags=re.UNICODE))


def _normalize_imported_message_html(raw_html: str) -> str:
    """Normalize ordinary imported content to editable Message defaults."""
    plain = _plain_text_from_html(raw_html).replace("\r\n", "\n").replace("\r", "\n")
    lines = plain.split("\n")
    blocks: list[str] = []
    for line in lines:
        if line.strip():
            blocks.append(f"<p>{_html.escape(line)}</p>")
        else:
            blocks.append("<p><br></p>")
    if not blocks:
        blocks = ["<p><br></p>"]
    return (
        '<div class="lettersmith-defaults" '
        "style=\"font-family:'Papyrus'; text-align:center; line-height:2;\">"
        + "".join(blocks)
        + "</div>"
    )


def _reading_time_label(word_count: int) -> str:
    if word_count <= 0:
        return "0 min read"
    if word_count < 200:
        return "<1 min read"
    return f"~{math.ceil(word_count / 200)} min read"


def _normalize_published_page_url(value: str) -> str:
    return normalize_published_page_url(
        value
    )


MESSAGE_OVERLAY_PRESET_KEY = "message_overlay_preset"
MESSAGE_OVERLAY_OPACITY_KEY = "message_overlay_opacity"
DEFAULT_MESSAGE_OVERLAY_PRESET = "paper"
DEFAULT_MESSAGE_OVERLAY_OPACITY = 68
TRANSPARENT_MESSAGE_SURFACE_OPACITY = 18
MESSAGE_OVERLAY_PRESETS: dict[str, tuple[tuple[int, int, int], str]] = {
    "black": ((0, 0, 0), "#ffffff"),
    "white": ((255, 255, 255), "#221710"),
    "paper": ((245, 235, 210), "#221710"),
    "clear": ((255, 255, 255), "#221710"),
}

MESSAGE_OVERLAY_PRESET_LABELS: dict[str, str] = {
    "paper": "Warm Paper",
    "black": "Dark Panel",
    "white": "Light Panel",
    "clear": "Transparent",
}


def _message_overlay_settings(settings_path: str | os.PathLike) -> tuple[str, int, tuple[int, int, int], str]:
    try:
        data = json.loads(Path(settings_path).read_text(encoding="utf-8")) if Path(settings_path).exists() else {}
    except Exception:
        data = {}

    preset = str(data.get(MESSAGE_OVERLAY_PRESET_KEY, DEFAULT_MESSAGE_OVERLAY_PRESET)).strip().lower()
    if preset not in MESSAGE_OVERLAY_PRESETS:
        preset = DEFAULT_MESSAGE_OVERLAY_PRESET

    try:
        opacity = int(data.get(MESSAGE_OVERLAY_OPACITY_KEY, DEFAULT_MESSAGE_OVERLAY_OPACITY))
    except Exception:
        opacity = DEFAULT_MESSAGE_OVERLAY_OPACITY
    opacity = max(0, min(100, opacity))

    if preset == "clear":
        opacity = 0

    rgb, ink = MESSAGE_OVERLAY_PRESETS[preset]
    return preset, opacity, rgb, ink


def _effective_message_overlay_opacity(
    preset: str,
    opacity: int,
) -> int:
    if preset == "clear":
        return TRANSPARENT_MESSAGE_SURFACE_OPACITY
    return max(0, min(100, int(opacity)))


def _soft_blur_message_background(image: QImage) -> QImage:
    """Blur a render copy without modifying the selected wall asset."""
    if image.isNull():
        return image
    reduced = image.scaled(
        max(1, image.width() // 18),
        max(1, image.height() // 18),
        Qt.IgnoreAspectRatio,
        Qt.SmoothTransformation,
    )
    return reduced.scaled(
        image.size(),
        Qt.IgnoreAspectRatio,
        Qt.SmoothTransformation,
    )


# ─────────────────────────────────────────────────────────────────────────────
# Drag-drop button
# ─────────────────────────────────────────────────────────────────────────────

class DropMessageButton(ArtworkButton):
    file_dropped = Signal(str)

    def __init__(self, project_root: str | Path, label: str = "Import") -> None:
        super().__init__(label, project_root, "EButton.png")
        self.setFont(QFont("Segoe UI Semibold", 24, QFont.Bold))
        self.setMinimumWidth(158)
        self.setFixedHeight(66)
        self.setSizePolicy(
            QtWidgets.QSizePolicy.Expanding,
            QtWidgets.QSizePolicy.Fixed,
        )
        self.set_artwork_fill(True)
        self.set_artwork_stretch(True)
        self.setAcceptDrops(True)
        self.setAccessibleName("Import")
        self.setToolTip(
            "Import a .txt, .docx, .pdf, .odt, or saved Letter Smith .html message. "
            "You may also drag a supported file onto this button."
        )
        if not self.has_artwork:
            self.setStyleSheet(self._default_style())

    def _default_style(self) -> str:
        return (
            "QPushButton {background-color:#1c1e26; border:1px solid #00d0ff;"
            " border-radius:8px; padding:8px 14px; color:#e6e6e6; text-align:center;}"
            "QPushButton:hover {background-color:#00d0ff; color:#0e0f12;}"
        )

    def _glow_style(self) -> str:
        return (
            "QPushButton {background-color:#1c1e26; border:2px solid #00ffff;"
            " border-radius:8px; padding:8px; color:#ffffff;}"
        )

    def _is_supported_url(self, url: QtCore.QUrl) -> bool:
        if not url.isLocalFile():
            return False
        suffix = url.toLocalFile().lower()
        return suffix.endswith((".txt", ".docx", ".pdf", ".odt", ".html", ".htm"))

    def dragEnterEvent(self, event: QtGui.QDragEnterEvent) -> None:  # type: ignore[override]
        if event.mimeData().hasUrls() and any(self._is_supported_url(u) for u in event.mimeData().urls()):
            event.acceptProposedAction()
            if not self.has_artwork:
                self.setStyleSheet(self._glow_style())
            self.update()
        else:
            event.ignore()

    def dragMoveEvent(self, event: QtGui.QDragMoveEvent) -> None:  # type: ignore[override]
        event.acceptProposedAction()

    def dragLeaveEvent(self, event: QtGui.QDragLeaveEvent) -> None:  # type: ignore[override]
        if not self.has_artwork:
            self.setStyleSheet(self._default_style())
        self.update()

    def dropEvent(self, event: QtGui.QDropEvent) -> None:  # type: ignore[override]
        if not self.has_artwork:
            self.setStyleSheet(self._default_style())
        self.update()
        for url in event.mimeData().urls():
            if self._is_supported_url(url):
                self.file_dropped.emit(url.toLocalFile())
                break


# ─────────────────────────────────────────────────────────────────────────────
# Revision history
# ─────────────────────────────────────────────────────────────────────────────

class RevisionHistoryDialog(QtWidgets.QDialog):
    def __init__(self, message_tab: "MessageTab") -> None:
        super().__init__(message_tab)
        self.message_tab = message_tab
        self.setWindowTitle("Message Revision History")
        self.setModal(True)
        self.resize(760, 480)

        root = QtWidgets.QVBoxLayout(self)
        root.setContentsMargins(14, 14, 14, 14)
        root.setSpacing(10)

        splitter = QtWidgets.QSplitter(Qt.Horizontal, self)
        self.revision_list = QtWidgets.QListWidget(splitter)
        self.preview = QtWidgets.QTextBrowser(splitter)
        self.preview.setOpenExternalLinks(False)
        splitter.setSizes([290, 470])
        root.addWidget(splitter, 1)

        buttons = QtWidgets.QHBoxLayout()
        self.restore_btn = QtWidgets.QPushButton("Restore", self)
        self.delete_btn = QtWidgets.QPushButton("Delete", self)
        close_btn = QtWidgets.QPushButton("Close", self)
        buttons.addWidget(self.restore_btn)
        buttons.addWidget(self.delete_btn)
        buttons.addStretch(1)
        buttons.addWidget(close_btn)
        root.addLayout(buttons)

        self.revision_list.currentItemChanged.connect(self._show_selected)
        self.revision_list.itemDoubleClicked.connect(lambda _item: self._restore_selected())
        self.restore_btn.clicked.connect(self._restore_selected)
        self.delete_btn.clicked.connect(self._delete_selected)
        close_btn.clicked.connect(self.accept)

        self.setStyleSheet(
            "QDialog{background:#121318;}"
            "QListWidget,QTextBrowser{background:#0d0f14;color:#e6e6e6;border:1px solid #2c3440;border-radius:7px;}"
            "QPushButton{background:#1c1e26;color:#e6e6e6;border:1px solid #00d0ff;border-radius:6px;padding:7px 14px;}"
            "QPushButton:hover{background:#00d0ff;color:#0e0f12;}"
        )
        self.refresh()

    def refresh(self) -> None:
        self.revision_list.clear()
        for revision in list_revisions(self.message_tab._html_path()):
            item = QtWidgets.QListWidgetItem(revision.display_name)
            item.setData(Qt.UserRole, str(revision.path))
            item.setToolTip(str(revision.path))
            self.revision_list.addItem(item)

        has_items = self.revision_list.count() > 0
        self.restore_btn.setEnabled(has_items)
        self.delete_btn.setEnabled(has_items)
        if has_items:
            self.revision_list.setCurrentRow(0)
        else:
            self.preview.setPlainText("No revisions have been created yet.")

    def _selected_path(self) -> Optional[Path]:
        item = self.revision_list.currentItem()
        if item is None:
            return None
        value = item.data(Qt.UserRole)
        return Path(str(value)) if value else None

    def _show_selected(self, current: Optional[QtWidgets.QListWidgetItem], _previous=None) -> None:
        if current is None:
            self.preview.clear()
            return
        path = self._selected_path()
        if path is None or not path.is_file():
            self.preview.setPlainText("This revision is no longer available.")
            return
        try:
            raw = path.read_text(encoding="utf-8")
            self.preview.setPlainText(_plain_text_from_html(raw))
        except Exception as error:
            self.preview.setPlainText(f"Could not read this revision: {error}")

    def _restore_selected(self) -> None:
        path = self._selected_path()
        if path is None:
            return
        if self.message_tab.restore_message_revision(path):
            self.accept()

    def _delete_selected(self) -> None:
        path = self._selected_path()
        if path is None:
            return
        try:
            delete_revision(path)
        except Exception as error:
            QtWidgets.QMessageBox.warning(self, "Revision History", f"Could not delete revision:\n{error}")
            return
        self.refresh()


# ─────────────────────────────────────────────────────────────────────────────
# Main Message Tab
# ─────────────────────────────────────────────────────────────────────────────

class MessageTab(QtWidgets.QWidget):
    # Public contract used by Nexus/Over_Nexus — DO NOT REMOVE
    text_selected = Signal(str)
    preview_image = Signal(QPixmap)
    wall_preview = Signal(QPixmap)
    published_page_url_changed = Signal(str)
    project_changed = Signal()

    def __init__(
        self,
        project_root: str,
        *,
        project_state: ProjectStateController | None = None,
        project_paths: ProjectPathResolver | None = None,
    ) -> None:
        super().__init__()
        self.project_root = project_root
        self.settings_path = os.path.join(project_root, SETTINGS_FILE)
        self.settings_store = SettingsStore(project_root)
        self.project_state = project_state
        if self.project_state is None:
            self.project_state = ProjectStateController(project_root)
            self.project_state.initialize()
        self.project_save_service = ProjectSaveService(
            project_root,
            self.project_state,
            resolver=project_paths,
        )

        # Compatibility cache; disk remains authoritative.
        self.current_html: str = ""
        self._content_has_intentional_formatting = False

        self._load_settings()
        (
            self.overlay_preset,
            self.overlay_opacity,
            _overlay_rgb,
            _overlay_ink,
        ) = _message_overlay_settings(self.settings_path)
        self.overlay_buttons: dict[str, QtWidgets.QPushButton] = {}
        self._overlay_render_timer = QtCore.QTimer(self)
        self._overlay_render_timer.setSingleShot(True)
        self._overlay_render_timer.setInterval(180)
        self._overlay_render_timer.timeout.connect(self._render_overlay_preview)
        self._sync_state = self._capture_sync_state()
        self._tab_active = False

        layout = QtWidgets.QVBoxLayout(self)
        layout.setContentsMargins(20, 12, 20, 14)
        layout.setSpacing(0)

        self.message_content_shell = QtWidgets.QFrame(self)
        self.message_content_shell.setObjectName("messageContentShell")
        self.message_content_shell.setMaximumWidth(820)
        self.message_content_shell.setSizePolicy(
            QtWidgets.QSizePolicy.Expanding,
            QtWidgets.QSizePolicy.Maximum,
        )
        self.message_content_shell.setStyleSheet(
            "QFrame#messageContentShell{background:transparent;border:none;}"
        )
        shell = QtWidgets.QVBoxLayout(self.message_content_shell)
        shell.setContentsMargins(0, 0, 0, 0)
        shell.setSpacing(10)

        header = QtWidgets.QLabel("Your Letter’s Message")
        header.setFont(QFont("Papyrus", 14))
        header.setStyleSheet("color:#00d0ff;")
        header.setAlignment(Qt.AlignCenter)
        shell.addWidget(header)

        self.title_recipient_container = QtWidgets.QWidget(self.message_content_shell)
        # Compatibility alias used by Over_Nexus and older code.
        self.title_sister_container = self.title_recipient_container
        title_recipient_layout = QtWidgets.QFormLayout(self.title_recipient_container)
        title_recipient_layout.setContentsMargins(0, 0, 0, 0)
        title_recipient_layout.setHorizontalSpacing(10)
        title_recipient_layout.setVerticalSpacing(7)
        title_recipient_layout.setFieldGrowthPolicy(QtWidgets.QFormLayout.AllNonFixedFieldsGrow)

        self.title_input = IdentityLineEdit(self.settings.get("recipient_title", ""))
        self.title_input.setPlaceholderText("e.g. Letter Title")
        self.name_input = IdentityLineEdit(self.settings.get("recipient_name", ""))
        self.name_input.setPlaceholderText("Recipient name")
        self.url_input = IdentityLineEdit(
            str(self.settings.get(PUBLISHED_PAGE_URL_KEY, ""))
        )
        self.url_input.setPlaceholderText("https://your-published-letter-page")
        self.url_input.setToolTip(
            "Save the public page address for this letter. "
            "Forge uses this address for Open Letter."
        )

        self._configure_identity_field(
            self.title_input,
            "title",
        )
        self._configure_identity_field(
            self.name_input,
            "recipient",
        )
        self._configure_identity_field(
            self.url_input,
            "published_url",
        )

        title_recipient_layout.addRow("Letter Title:", self.title_input)
        title_recipient_layout.addRow("Recipient:", self.name_input)
        title_recipient_layout.addRow("Published Page URL:", self.url_input)

        self.title_input.editingFinished.connect(self._save_settings)
        self.name_input.editingFinished.connect(self._save_settings)
        self.url_input.editingFinished.connect(self._save_settings)
        shell.addWidget(self.title_recipient_container)

        shell.addWidget(self._build_message_overlay_controls())

        actions = QtWidgets.QGridLayout()
        actions.setContentsMargins(0, 0, 0, 0)
        actions.setHorizontalSpacing(10)
        actions.setVerticalSpacing(8)
        actions.setColumnStretch(0, 13)
        actions.setColumnStretch(1, 8)
        actions.setColumnStretch(2, 8)
        actions.setColumnStretch(3, 8)
        actions.setColumnStretch(4, 13)

        compact_button_style = (
            "QPushButton{min-height:40px;padding:0 16px;background:#171b20;color:#e4ebf4;"
            "border:1px solid #38424f;border-radius:7px;font-weight:700;}"
            "QPushButton:hover{border-color:#00d0ff;background:#1c252e;}"
            "QPushButton:disabled{color:#69727e;border-color:#2b3139;background:#15181c;}"
        )

        self.btn = DropMessageButton(self.project_root, "Import")
        self.btn.clicked.connect(self.select_file)
        self.btn.file_dropped.connect(self.handle_drop)
        actions.addWidget(self.btn, 0, 1)

        self.edit_btn = ArtworkButton("Edit", self.project_root, "ROButton.png", self)
        self.edit_btn.setMinimumWidth(158)
        self.edit_btn.setFixedHeight(66)
        self.edit_btn.setSizePolicy(
            QtWidgets.QSizePolicy.Expanding,
            QtWidgets.QSizePolicy.Fixed,
        )
        self.edit_btn.set_artwork_fill(True)
        self.edit_btn.set_artwork_stretch(True)
        self.edit_btn.setFont(QFont("Segoe UI Semibold", 24, QFont.Bold))
        if not self.edit_btn.has_artwork:
            self.edit_btn.setStyleSheet(compact_button_style)
        self.edit_btn.setToolTip("Open the rich-text editor for the current message.")
        self.edit_btn.setAccessibleName("Edit")
        self.edit_btn.setEnabled(True)
        self.edit_btn.clicked.connect(self.open_editor)
        actions.addWidget(self.edit_btn, 0, 2)

        self.revisions_btn = ArtworkButton("Revisions", self.project_root, "RButton.png", self)
        self.revisions_btn.setMinimumWidth(158)
        self.revisions_btn.setFixedHeight(66)
        self.revisions_btn.setSizePolicy(
            QtWidgets.QSizePolicy.Expanding,
            QtWidgets.QSizePolicy.Fixed,
        )
        self.revisions_btn.set_artwork_fill(True)
        self.revisions_btn.set_artwork_stretch(True)
        self.revisions_btn.setFont(QFont("Segoe UI Semibold", 24, QFont.Bold))
        if not self.revisions_btn.has_artwork:
            self.revisions_btn.setStyleSheet(compact_button_style)
        self.revisions_btn.setToolTip("Open autosaved message versions and restore an earlier version.")
        self.revisions_btn.setAccessibleName("Revisions")
        self.revisions_btn.clicked.connect(self.open_revision_history)
        actions.addWidget(self.revisions_btn, 0, 3)
        shell.addLayout(actions)

        self.status = QtWidgets.QLabel()
        self.status.setFont(QFont("Segoe UI", 9))
        self.status.setAlignment(Qt.AlignCenter)
        self.status.setStyleSheet("color:#aeb8c6; min-height:16px;")
        shell.addWidget(self.status)

        self.message_summary = QtWidgets.QLabel("0 words  •  0 characters  •  0 min read")
        self.message_summary.setAlignment(Qt.AlignCenter)
        self.message_summary.setStyleSheet(
            "color:#8794a5; font:9px 'Segoe UI'; padding-top:1px;"
        )
        shell.addWidget(self.message_summary)

        layout.addWidget(self.message_content_shell, 0, Qt.AlignHCenter | Qt.AlignTop)
        layout.addStretch(1)

        # Ensure fallback assets and load the current message immediately.
        self._check_existing()
        self._update_message_summary()

    # ──────────────────────────────────────────────────────────────────
    # Show hook: every time user clicks into Message tab
    # ──────────────────────────────────────────────────────────────────
    def _message_settings_signature(self) -> str:
        settings = self.settings_store.snapshot()
        relevant = {
            key: settings.get(key)
            for key in (
                "recipient_title",
                "recipient_name",
                PUBLISHED_PAGE_URL_KEY,
                MESSAGE_OVERLAY_PRESET_KEY,
                MESSAGE_OVERLAY_OPACITY_KEY,
            )
        }
        return json.dumps(relevant, sort_keys=True, separators=(",", ":"), ensure_ascii=False)

    def _capture_sync_state(self) -> dict[str, str]:
        return {
            "html": file_fingerprint(self._html_path()),
            "png": file_fingerprint(self._png_path()),
            "wall": file_fingerprint(self._wall_path()),
            "settings": self._message_settings_signature(),
        }

    def sync_from_disk(self, *, force: bool = False) -> bool:
        """Reconcile Message content, metadata, rendering, and direct disk edits."""
        before = dict(getattr(self, "_sync_state", {}))
        after = self._capture_sync_state()
        changed_keys = {key for key, value in after.items() if force or before.get(key) != value}

        self._load_settings()
        self.title_input.setText(str(self.settings.get("recipient_title", "")))
        self.name_input.setText(str(self.settings.get("recipient_name", "")))
        self.url_input.setText(str(self.settings.get(PUBLISHED_PAGE_URL_KEY, "")))
        self._sync_identity_field_lock(self.title_input, "title")
        self._sync_identity_field_lock(self.name_input, "recipient")
        self._sync_identity_field_lock(self.url_input, "published_url")
        (
            self.overlay_preset,
            self.overlay_opacity,
            _overlay_rgb,
            _overlay_ink,
        ) = _message_overlay_settings(self.settings_path)
        self._sync_overlay_controls()
        self._refresh_message_from_disk()
        self._ensure_wall_exists()

        render_changed = bool(changed_keys.intersection({"html", "wall", "settings"}))
        if render_changed and self.current_html.strip():
            self._generate_image(self.current_html)
        else:
            self._ensure_message_exists()

        self._emit_best_preview()
        self._update_message_summary()
        self._sync_state = self._capture_sync_state()
        if changed_keys:
            self.project_changed.emit()
        return bool(changed_keys)

    def sync_to_disk(self) -> bool:
        """Persist Message metadata and finish render state before tab exit."""
        before = dict(
            getattr(
                self,
                "_sync_state",
                {},
            )
        )
        disk_state = self._capture_sync_state()
        render_changed = any(
            before.get(key)
            != disk_state.get(key)
            for key in (
                "html",
                "wall",
                "settings",
            )
        )
        self._save_settings()
        self._refresh_message_from_disk()
        if self.current_html.strip():
            try:
                self.project_save_service.save_message(
                    self.current_html,
                    workspace_path=self._html_path(),
                    reason="message-tab-exit",
                )
            except ProjectNotReadyError:
                pass
        self._ensure_wall_exists()
        if (
            self.current_html.strip()
            and (
                render_changed
                or not self._png_path().is_file()
            )
        ):
            self._generate_image(self.current_html)
            self._sync_project_render_assets()
        self._ensure_message_exists()
        self._sync_state = self._capture_sync_state()
        changed = before != self._sync_state
        if changed:
            self.project_changed.emit()
        return changed

    def activate_for_tab_change(self) -> None:
        if self._tab_active:
            return
        self._tab_active = True
        self.sync_from_disk()

    def deactivate_for_tab_change(self) -> None:
        if not self._tab_active:
            return
        self.sync_to_disk()
        self._tab_active = False

    def showEvent(self, event: QtGui.QShowEvent) -> None:  # type: ignore[override]
        super().showEvent(event)
        self.activate_for_tab_change()

    def hideEvent(self, event: QtGui.QHideEvent) -> None:  # type: ignore[override]
        self.deactivate_for_tab_change()
        super().hideEvent(event)

    def refresh_from_disk(self) -> None:
        self.sync_from_disk(force=True)

    def focus_field(self, target: str) -> None:
        """Focus the Message correction requested by Project Readiness."""
        widget = {
            "recipient": self.name_input,
            "title": self.title_input,
            "published_url": self.url_input,
            "message": self.edit_btn,
        }.get(str(target))
        if widget is not None:
            widget.setFocus(Qt.OtherFocusReason)

    # ──────────────────────────────────────────────────────────────────
    # Nexus / Over_Nexus hooks
    # ──────────────────────────────────────────────────────────────────
    def toggle_title_recipient_area(self) -> None:
        self.title_recipient_container.setVisible(not self.title_recipient_container.isVisible())

    def toggle_title_sister_area(self) -> None:
        # Compatibility shim for older Nexus wiring.
        self.toggle_title_recipient_area()

    def open_message_editor(self) -> None:
        self.open_editor()

    # ──────────────────────────────────────────────────────────────────
    # Message overlay controls
    # ──────────────────────────────────────────────────────────────────
    def _build_message_overlay_controls(self) -> QtWidgets.QFrame:
        panel = QtWidgets.QFrame(self.message_content_shell)
        panel.setObjectName("messageOverlayControls")
        panel.setMaximumHeight(86)
        panel.setStyleSheet(
            "QFrame#messageOverlayControls{background:#15191f;border:1px solid #303945;border-radius:8px;}"
            "QLabel{color:#cfd8e5;background:transparent;}"
            "QComboBox{min-height:30px;padding:0 28px 0 9px;background:#1d232b;color:#eef4fb;"
            "border:1px solid #435064;border-radius:6px;}"
            "QComboBox:hover{border-color:#00d0ff;}"
            "QComboBox::drop-down{width:24px;border:none;}"
            "QComboBox QAbstractItemView{background:#151a21;color:#eef4fb;border:1px solid #435064;"
            "selection-background-color:#254252;outline:0;}"
            "QSlider::groove:horizontal{height:4px;background:#272e38;border-radius:2px;}"
            "QSlider::sub-page:horizontal{background:#447b8a;border-radius:2px;}"
            "QSlider::handle:horizontal{width:13px;margin:-5px 0;background:#d9e4ef;border-radius:6px;}"
            "QSlider:disabled::handle:horizontal{background:#68717c;}"
        )

        root = QtWidgets.QGridLayout(panel)
        root.setContentsMargins(10, 8, 10, 8)
        root.setHorizontalSpacing(10)
        root.setVerticalSpacing(5)
        root.setColumnStretch(1, 1)

        title = QtWidgets.QLabel("Text background", panel)
        title.setStyleSheet("font-weight:700;color:#eef4fb;")
        root.addWidget(title, 0, 0)

        self.overlay_preset_combo = QtWidgets.QComboBox(panel)
        self.overlay_preset_combo.setFixedWidth(210)
        for key in ("paper", "black", "white", "clear"):
            self.overlay_preset_combo.addItem(MESSAGE_OVERLAY_PRESET_LABELS[key], key)
        self.overlay_preset_combo.currentIndexChanged.connect(self._on_overlay_preset_changed)
        root.addWidget(self.overlay_preset_combo, 0, 1)

        self.overlay_opacity_label = QtWidgets.QLabel(panel)
        self.overlay_opacity_label.setFixedWidth(112)
        self.overlay_opacity_label.setStyleSheet("color:#aeb8c6;font:9px 'Segoe UI';")
        root.addWidget(self.overlay_opacity_label, 1, 0)

        self.overlay_opacity_slider = QtWidgets.QSlider(Qt.Horizontal, panel)
        self.overlay_opacity_slider.setRange(0, 100)
        self.overlay_opacity_slider.setValue(int(self.overlay_opacity))
        self.overlay_opacity_slider.setMaximumWidth(260)
        self.overlay_opacity_slider.valueChanged.connect(self._set_overlay_opacity)
        root.addWidget(self.overlay_opacity_slider, 1, 1)

        self._sync_overlay_controls()
        return panel

    def _on_overlay_preset_changed(self, index: int) -> None:
        preset = self.overlay_preset_combo.itemData(index)
        self._set_overlay_preset(str(preset or ""))

    def _set_overlay_preset(self, preset: str) -> None:
        preset = str(preset or "").strip().lower()
        if preset not in MESSAGE_OVERLAY_PRESETS:
            return
        self.overlay_preset = preset
        if preset == "clear":
            self.overlay_opacity = 0
        elif self.overlay_opacity <= 0:
            self.overlay_opacity = DEFAULT_MESSAGE_OVERLAY_OPACITY
        self._persist_overlay_settings()

    def _set_overlay_opacity(self, value: int) -> None:
        self.overlay_opacity = max(0, min(100, int(value)))
        if self.overlay_preset == "clear" and self.overlay_opacity > 0:
            self.overlay_preset = DEFAULT_MESSAGE_OVERLAY_PRESET
        self._persist_overlay_settings()

    def _persist_overlay_settings(self) -> None:
        self.settings[MESSAGE_OVERLAY_PRESET_KEY] = self.overlay_preset
        self.settings[MESSAGE_OVERLAY_OPACITY_KEY] = int(self.overlay_opacity)
        self._persist_settings(announce=False)
        self._sync_overlay_controls()
        self._overlay_render_timer.start()

    def _render_overlay_preview(self) -> None:
        html = self.current_html or "<p><br></p>"
        self._generate_image(html)
        self._emit_best_preview()

    def _sync_overlay_controls(self) -> None:
        if hasattr(self, "overlay_preset_combo"):
            target_index = self.overlay_preset_combo.findData(self.overlay_preset)
            self.overlay_preset_combo.blockSignals(True)
            if target_index >= 0:
                self.overlay_preset_combo.setCurrentIndex(target_index)
            self.overlay_preset_combo.blockSignals(False)

        is_transparent = self.overlay_preset == "clear"
        if hasattr(self, "overlay_opacity_slider"):
            self.overlay_opacity_slider.blockSignals(True)
            self.overlay_opacity_slider.setValue(int(self.overlay_opacity))
            self.overlay_opacity_slider.setEnabled(not is_transparent)
            self.overlay_opacity_slider.blockSignals(False)

        if hasattr(self, "overlay_opacity_label"):
            if is_transparent:
                self.overlay_opacity_label.setText("Blurred glass")
            else:
                self.overlay_opacity_label.setText(f"Opacity: {int(self.overlay_opacity)}%")


    # ──────────────────────────────────────────────────────────────────
    # Summary + revision history
    # ──────────────────────────────────────────────────────────────────
    def _update_message_summary(self, html: Optional[str] = None) -> None:
        raw = self.current_html if html is None else html
        plain = _plain_text_from_html(raw or "")
        words = _word_count(plain)
        characters = len(plain)
        self.message_summary.setText(
            f"{words:,} words  •  {characters:,} characters  •  {_reading_time_label(words)}"
        )

    def _refresh_message_from_disk(self) -> None:
        path = self._html_path()
        if not path.is_file():
            self.current_html = ""
            self._content_has_intentional_formatting = False
            return
        try:
            self.current_html = path.read_text(encoding="utf-8")
            self._content_has_intentional_formatting = True
        except Exception:
            pass

    def open_revision_history(self) -> None:
        RevisionHistoryDialog(self).exec()

    def restore_message_revision(self, revision_path: Path) -> bool:
        try:
            restored = restore_revision(self._html_path(), revision_path)
        except Exception as error:
            QtWidgets.QMessageBox.warning(self, "Revision History", f"Could not restore revision:\n{error}")
            return False

        self.current_html = restored
        try:
            self.project_save_service.save_message(
                restored,
                workspace_path=self._html_path(),
                reason="revision-restore",
            )
        except Exception as error:
            self.status.setText(
                f"Could not save restored revision: {error}"
            )
            return False
        self._content_has_intentional_formatting = True
        self.text_selected.emit(restored)
        self._update_message_summary(restored)
        self._ensure_wall_exists()
        self._generate_image(restored)
        self._sync_project_render_assets()
        self._emit_best_preview()
        self.status.setText("Revision restored.")
        self._sync_state = self._capture_sync_state()
        self.project_changed.emit()
        return True

    # ──────────────────────────────────────────────────────────────────
    # Settings
    # ──────────────────────────────────────────────────────────────────
    def _load_settings(self) -> None:
        self.settings = self.settings_store.snapshot()

    @staticmethod
    def _setting_bool(value: object) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            return value.strip().casefold() in {"1", "true", "yes", "on"}
        return bool(value)

    def _configure_identity_field(
        self,
        field: IdentityLineEdit,
        field_name: str,
    ) -> None:
        lock_key = IDENTITY_LOCK_KEYS[field_name]
        field.setReadOnly(
            self._setting_bool(self.settings.get(lock_key, False))
            and bool(field.text().strip())
        )
        field.double_clicked.connect(
            lambda field=field, field_name=field_name: self._unlock_identity_field(
                field,
                field_name,
            )
        )
        field.returnPressed.connect(
            lambda field=field, field_name=field_name: self._commit_identity_field(
                field,
                field_name,
            )
        )
        self._style_identity_field(field)

    def _sync_identity_field_lock(
        self,
        field: IdentityLineEdit,
        field_name: str,
    ) -> None:
        field.setReadOnly(
            self._setting_bool(
                self.settings.get(IDENTITY_LOCK_KEYS[field_name], False)
            )
            and bool(field.text().strip())
        )
        self._style_identity_field(field)

    @staticmethod
    def _style_identity_field(field: IdentityLineEdit) -> None:
        if field.isReadOnly():
            field.setStyleSheet(
                "QLineEdit{background:#10263b;color:#78b9dd;"
                "border:1px solid #244b69;border-radius:5px;padding:5px;}"
                "QLineEdit:focus{border-color:#3d8fba;}"
                "QLineEdit:selected{background:#1c5275;color:#e8f7ff;}"
            )
            field.setToolTip(
                "Committed. Double-click to edit, then press Enter to commit again."
            )
        else:
            field.setStyleSheet(
                "QLineEdit{background:#121b23;color:#e8f9ff;"
                "border:1px solid #38424f;border-radius:5px;padding:5px;}"
                "QLineEdit:focus{border-color:#00d2ef;}"
            )

    def _unlock_identity_field(
        self,
        field: IdentityLineEdit,
        field_name: str,
    ) -> None:
        field.setReadOnly(False)
        self.settings[IDENTITY_LOCK_KEYS[field_name]] = False
        self._style_identity_field(field)
        field.setFocus(Qt.MouseFocusReason)

    def _commit_identity_field(
        self,
        field: IdentityLineEdit,
        field_name: str,
    ) -> None:
        value = field.text().strip()
        if not value:
            return
        if field_name == "published_url" and not _normalize_published_page_url(value):
            self.status.setText(
                "Published Page URL must be a valid HTTP or HTTPS address."
            )
            return
        if not self._save_settings():
            return
        field.setReadOnly(True)
        self.settings[IDENTITY_LOCK_KEYS[field_name]] = True
        if self._persist_settings(announce=False):
            self._style_identity_field(field)

    def reset_identity_locks(self) -> None:
        """Clear identity fields and leave all three controls editable."""
        for field_name, lock_key in IDENTITY_LOCK_KEYS.items():
            field = {
                "title": self.title_input,
                "recipient": self.name_input,
                "published_url": self.url_input,
            }[field_name]
            field.setReadOnly(False)
            self.settings[lock_key] = False
            self._style_identity_field(field)

    def _persist_settings(self, *, announce: bool) -> bool:
        if not self.project_state.is_project_ready:
            if announce:
                self.status.setText(
                    "A recipient is required before saving."
                )
            return False
        try:
            fields = {
                key: self.settings[key]
                for key in (
                    "recipient_title",
                    "recipient_name",
                    PUBLISHED_PAGE_URL_KEY,
                    MESSAGE_OVERLAY_PRESET_KEY,
                    MESSAGE_OVERLAY_OPACITY_KEY,
                    *IDENTITY_LOCK_KEYS.values(),
                )
                if key in self.settings
            }
            self.settings = self.settings_store.update_fields(fields)
            if announce:
                self.status.setText("Message details saved.")
        except Exception as error:
            if announce:
                self.status.setText(f"Error saving settings: {error}")
            return False
        self.project_changed.emit()
        return True

    def _save_settings(self) -> bool:
        self.settings["recipient_title"] = self.title_input.text().strip()
        self.settings["recipient_name"] = self.name_input.text().strip()

        raw_url = self.url_input.text().strip()
        normalized_url = _normalize_published_page_url(raw_url)
        if raw_url and not normalized_url:
            # Preserve the last valid saved URL instead of replacing it with an
            # unusable value. Title and recipient changes are still saved.
            self._persist_settings(announce=False)
            self.status.setText(
                "Title and recipient saved. Published Page URL must be a valid HTTP or HTTPS address."
            )
            return False

        self.settings[PUBLISHED_PAGE_URL_KEY] = normalized_url
        self.url_input.setText(normalized_url)
        if self._persist_settings(announce=True):
            self.published_page_url_changed.emit(normalized_url)
            return True
        return False

    def set_published_page_url(self, url: str, *, persist: bool = True, announce: bool = True) -> bool:
        raw_url = (url or "").strip()
        normalized_url = _normalize_published_page_url(raw_url)
        if raw_url and not normalized_url:
            if announce:
                self.status.setText("Published Page URL must be a valid HTTP or HTTPS address.")
            return False

        self.settings[PUBLISHED_PAGE_URL_KEY] = normalized_url
        if hasattr(self, "url_input"):
            self.url_input.setText(normalized_url)
        if persist:
            if not self._persist_settings(announce=announce):
                return False
        else:
            self.settings = self.settings_store.snapshot()
            self.settings[PUBLISHED_PAGE_URL_KEY] = normalized_url
        self.published_page_url_changed.emit(normalized_url)
        return True

    # ──────────────────────────────────────────────────────────────────
    # Paths (AUTHORITATIVE)
    # ──────────────────────────────────────────────────────────────────
    def _html_path(self) -> Path:
        # SOURCE OF TRUTH: gallery/user/message/message.html
        return Path(self.project_root) / MESSAGE_HTML_FILE

    def _png_path(self) -> Path:
        # SOURCE OF TRUTH: gallery/user/message/message.png
        return Path(self.project_root) / MESSAGE_IMAGE_FILE

    def _wall_path(self) -> Path:
        # SOURCE OF TRUTH: gallery/user/pages/wall.png
        return Path(self.project_root) / USER_PAGES_DIR / "wall.png"

    def _sync_project_render_assets(self) -> None:
        if not self.project_state.is_project_ready:
            return
        for source, relative in (
            (
                self._png_path(),
                Path("message") / "message.png",
            ),
            (
                self._wall_path(),
                Path("pages") / "wall.png",
            ),
        ):
            if not source.is_file():
                continue
            try:
                self.project_save_service.copy_workspace_file(
                    source,
                    relative,
                )
            except Exception as error:
                self.status.setText(
                    f"Could not save project artwork: {error}"
                )

    def _default_message_bg_path(self) -> Path:
        # DEFAULT FALLBACK: gallery/app/pages/Dmessage.png
        return Path(self.project_root) / "gallery" / "app" / "pages" / "Dmessage.png"

    # ──────────────────────────────────────────────────────────────────
    # Fallback pipeline (Dmessage → wall, then wall → message.png if needed)
    # ──────────────────────────────────────────────────────────────────
    def _ensure_wall_exists(self) -> bool:
        """
        Guarantees wall.png exists by copying:
            gallery/app/pages/Dmessage.png -> gallery/user/pages/wall.png
        if wall.png is missing.
        """
        wall_path = self._wall_path()
        if wall_path.exists():
            return True

        src = self._default_message_bg_path()
        try:
            wall_path.parent.mkdir(parents=True, exist_ok=True)
            if src.exists():
                shutil.copyfile(str(src), str(wall_path))
                return True
        except Exception:
            pass
        return wall_path.exists()

    def _ensure_message_exists(self) -> None:
        """
        Guarantees message.png exists.
        - If missing, render using current_html if available,
          else render a blank message (wall-only).
        - If render fails, force message.png = wall.png (scaled).
        """
        out_png = self._png_path()
        if out_png.exists():
            return

        # Ensure wall is there first
        self._ensure_wall_exists()

        # Try render from HTML if any, else blank
        try:
            html = (self._html_path().read_text(encoding="utf-8") if self._html_path().exists() else "").strip()
            if not html:
                html = "<p></p>"
            self._generate_image(html)
            if out_png.exists():
                return
        except Exception:
            pass

        # Hard fallback: copy wall into message.png (scaled to 2048×3072)
        try:
            wall_path = self._wall_path()
            if wall_path.exists():
                wall_img = QImage(str(wall_path))
                if not wall_img.isNull():
                    FULL_W, FULL_H = 2048, 3072
                    wall_img = wall_img.scaled(
                        FULL_W, FULL_H,
                        Qt.KeepAspectRatioByExpanding,
                        Qt.SmoothTransformation
                    )
                    canvas = QImage(FULL_W, FULL_H, QImage.Format_ARGB32)
                    canvas.fill(QtGui.QColor(14, 14, 18, 255))
                    p = QPainter(canvas)
                    p.setRenderHint(QPainter.SmoothPixmapTransform, True)
                    p.drawImage(0, 0, wall_img)
                    p.end()
                    _atomic_save_image(canvas, str(out_png))
        except Exception:
            pass

    # ──────────────────────────────────────────────────────────────────
    # Existing content load
    # ──────────────────────────────────────────────────────────────────
    def _check_existing(self) -> None:
        self._ensure_wall_exists()
        self._ensure_message_exists()

        html_path = self._html_path()
        if html_path.is_file():
            try:
                self.current_html = html_path.read_text(encoding="utf-8")
                # Canonical message.html is a saved Letter Smith message and must
                # retain its intentional user formatting.
                self._content_has_intentional_formatting = True
                self.text_selected.emit(self.current_html)
            except Exception as error:
                self.status.setText(f"Failed to read message.html: {error}")
        else:
            self.current_html = ""
            self._content_has_intentional_formatting = False

        self.edit_btn.setEnabled(True)
        self._update_message_summary()
        self._emit_best_preview()

    # ──────────────────────────────────────────────────────────────────
    # Preview helpers
    # ──────────────────────────────────────────────────────────────────
    def _emit_best_preview(self) -> None:
        """
        Rule:
        - If message.png exists, show it.
        - Else show wall.png (which is guaranteed by Dmessage fallback).
        """
        png_path = self._png_path()
        if png_path.is_file():
            full_pix = QPixmap(str(png_path))
            if not full_pix.isNull():
                thumb = full_pix.scaled(169, 253, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                self.preview_image.emit(thumb)
                self.status.setText("🖼️ Showing message.png")
                return

        wall_path = self._wall_path()
        if wall_path.is_file():
            full_pix = QPixmap(str(wall_path))
            if not full_pix.isNull():
                thumb = full_pix.scaled(169, 253, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                self.preview_image.emit(thumb)
                self.status.setText("🧱 Showing wall.png (fallback)")
                return

        # If we get here, fallback assets are missing/corrupt
        self.status.setText("❌ No message.png or wall.png available (fallback missing).")

    # ──────────────────────────────────────────────────────────────────
    # File selection / drop
    # ──────────────────────────────────────────────────────────────────
    def select_file(self) -> None:
        path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self,
            "Select Message File",
            "",
            "Messages (*.txt *.docx *.pdf *.odt *.html *.htm)",
        )
        if path:
            self._process_file(path)

    def handle_drop(self, path: str) -> None:
        if os.path.exists(path):
            self._process_file(path)

    # ──────────────────────────────────────────────────────────────────
    # Editor
    # ──────────────────────────────────────────────────────────────────
    def open_editor(self) -> None:
        if not self.project_state.is_project_ready:
            self.status.setText(
                "A recipient is required before editing."
            )
            return
        html_path = self._html_path()
        if html_path.is_file():
            try:
                html_for_editor = html_path.read_text(encoding="utf-8")
            except Exception:
                html_for_editor = self.current_html or ""
        else:
            html_for_editor = self.current_html or ""

        full_pix: Optional[QPixmap] = None
        png_path = self._png_path()
        if png_path.is_file():
            candidate = QPixmap(str(png_path))
            if not candidate.isNull():
                full_pix = candidate

        dlg = Editor(
            html_for_editor,
            full_pix,
            parent=self,
            apply_defaults=not self._content_has_intentional_formatting,
        )
        dlg.autosaved.connect(self._handle_editor_autosaved)
        dlg.finished.connect(lambda _result: self._handle_editor_finished(dlg))
        dlg.exec()

    def _handle_editor_autosaved(self, html: str) -> None:
        if not html:
            return
        self.current_html = html
        self._content_has_intentional_formatting = True
        self._update_message_summary(html)
        self.text_selected.emit(html)
        self._sync_state = self._capture_sync_state()
        self.project_changed.emit()

    def _handle_editor_finished(self, dlg: QtWidgets.QDialog) -> None:
        try:
            new_html = dlg.get_edited_html()  # type: ignore[attr-defined]
        except Exception:
            new_html = self.current_html

        if not new_html:
            return

        self.current_html = new_html
        self._content_has_intentional_formatting = True
        self._update_message_summary(new_html)
        self._ensure_wall_exists()
        self.text_selected.emit(new_html)
        self._generate_image(new_html)
        self._sync_project_render_assets()
        self._emit_best_preview()
        self.status.setText("Message saved.")
        self.project_changed.emit()

    # ──────────────────────────────────────────────────────────────────
    # Preview button
    # ──────────────────────────────────────────────────────────────────
    def _emit_preview(self) -> None:
        # Preview button should prefer message.png; if missing, show wall fallback.
        self._ensure_wall_exists()
        self._ensure_message_exists()
        self._emit_best_preview()

    # ──────────────────────────────────────────────────────────────────
    # Extraction + processing
    # ──────────────────────────────────────────────────────────────────
    def extract_text(self, path: str) -> str:
        low = path.lower()
        try:
            if low.endswith((".html", ".htm")):
                return Path(path).read_text(encoding="utf-8", errors="ignore")

            if low.endswith(".txt"):
                with open(path, "r", encoding="utf-8") as f:
                    raw = f.read()
                esc = (
                    raw.replace("&", "&amp;")
                       .replace("<", "&lt;")
                       .replace(">", "&gt;")
                )
                return esc.replace("\n", "<br>")

            if low.endswith(".docx"):
                if mammoth is not None:
                    try:
                        with open(path, "rb") as docx_file:
                            res = mammoth.convert_to_html(docx_file)
                        return res.value
                    except Exception:
                        pass
                if Document is not None:
                    try:
                        doc = Document(path)
                        parts = []
                        for p in doc.paragraphs:
                            t = p.text.strip()
                            if t:
                                parts.append(_html.escape(t))
                        return "<br>".join(parts)
                    except Exception:
                        pass
                return ""

            if low.endswith(".pdf"):
                if PdfReader is not None:
                    try:
                        reader = PdfReader(path)
                        pages = [(page.extract_text() or "").replace("\n", "<br>") for page in reader.pages]
                        return "<br><br>".join(pages)
                    except Exception:
                        pass
                try:
                    import pdfplumber  # type: ignore
                    txt_pages = []
                    with pdfplumber.open(path) as pdf:
                        for page in pdf.pages:
                            txt = (page.extract_text() or "").strip()
                            if txt:
                                txt_pages.append(_html.escape(txt).replace("\n", "<br>"))
                    return "<br><br>".join(txt_pages)
                except Exception:
                    return ""

            if low.endswith(".odt"):
                try:
                    from odf.opendocument import load as _load  # type: ignore
                    from odf.text import P as _P  # type: ignore
                    doc = _load(path)
                    paras = doc.getElementsByType(_P)
                    chunks = []
                    for p in paras:
                        text = "".join(getattr(n, "data", "") for n in p.childNodes)
                        text = _html.escape(text)
                        if text.strip():
                            chunks.append(text)
                    return "<br>".join(chunks)
                except Exception:
                    return ""
        except Exception as e:
            self.status.setText(f"❌ Error reading file: {e}")
        return ""

    def _process_file(self, path: str) -> None:
        imported_html = self.extract_text(path)
        if not imported_html.strip():
            self.status.setText("That file had no extractable text.")
            return

        preserve_formatting = is_lettersmith_message_html(imported_html, filename=path)
        html = imported_html if preserve_formatting else _normalize_imported_message_html(imported_html)

        try:
            self.project_save_service.save_message(
                html,
                workspace_path=self._html_path(),
                reason="import",
            )
            self.current_html = html
            self._content_has_intentional_formatting = preserve_formatting
            self.edit_btn.setEnabled(True)
            self.status.setText(f"Message imported from {Path(path).name}.")
            self.text_selected.emit(html)
        except Exception as error:
            self.status.setText(f"Error saving message.html: {error}")
            return

        self._update_message_summary(html)
        self._ensure_wall_exists()
        self._generate_image(html)
        self._sync_project_render_assets()
        self._emit_best_preview()
        self.project_changed.emit()

    # ──────────────────────────────────────────────────────────────────
    # Render message.png (stable, saved in same folder as message.html)
    # ──────────────────────────────────────────────────────────────────
    def _generate_image(self, html: str) -> None:
        """
        Produces (SOURCE):
            gallery/user/message/message.png

        Fallback rules:
        - wall.png is guaranteed (Dmessage -> wall if missing)
        - if save fails, we also ensure message.png exists via _ensure_message_exists()
        """
        try:
            # Guarantee wall exists first
            self._ensure_wall_exists()

            FULL_W, FULL_H = 2048, 3072
            MARGIN_LR = 100
            MARGIN_TOP = 100
            MARGIN_BOTTOM = 100
            TEXT_WIDTH = FULL_W - 2 * MARGIN_LR
            TEXT_HEIGHT = FULL_H - MARGIN_TOP - MARGIN_BOTTOM

            out_png = self._png_path()
            out_png.parent.mkdir(parents=True, exist_ok=True)

            # Opaque base first (prevents “transparent saved as black” surprises)
            canvas = QImage(FULL_W, FULL_H, QImage.Format_ARGB32)
            canvas.fill(QtGui.QColor(14, 14, 18, 255))

            painter = QPainter(canvas)
            painter.setRenderHint(QPainter.Antialiasing, True)
            painter.setRenderHint(QPainter.TextAntialiasing, True)
            painter.setRenderHint(QPainter.SmoothPixmapTransform, True)

            preset, overlay_opacity, overlay_rgb, ink_color = (
                _message_overlay_settings(self.settings_path)
            )

            # Background wall (SOURCE): gallery/user/pages/wall.png
            wall_path = self._wall_path()
            if wall_path.exists():
                wall_img = QImage(str(wall_path))
                if not wall_img.isNull():
                    wall_img = wall_img.scaled(
                        FULL_W,
                        FULL_H,
                        Qt.KeepAspectRatioByExpanding,
                        Qt.SmoothTransformation
                    )
                    if preset == "clear":
                        wall_img = _soft_blur_message_background(
                            wall_img
                        )
                    painter.drawImage(0, 0, wall_img)

            # User-controlled text background overlay.
            surface_opacity = _effective_message_overlay_opacity(
                preset,
                overlay_opacity,
            )
            if surface_opacity > 0:
                r, g, b = overlay_rgb
                overlay = QtGui.QColor(
                    r,
                    g,
                    b,
                    int(round(255 * (surface_opacity / 100.0))),
                )
                painter.fillRect(0, 0, FULL_W, FULL_H, overlay)

            # HTML text
            doc = QTextDocument()
            doc.setDefaultFont(QFont("Papyrus", 12))

            doc.setDefaultStyleSheet(
                f"body {{ color: {ink_color}; background: transparent; font-family: 'Papyrus'; text-align: center; line-height: 2; }}"
                "p { margin: 0 0 12px 0; }"
                "br { line-height: 2; }"
            )

            doc.setHtml(html)
            normalize_ultralinks_in_document(doc)
            doc.setTextWidth(TEXT_WIDTH)
            doc.setPageSize(QSizeF(TEXT_WIDTH, TEXT_HEIGHT))

            painter.save()
            painter.translate(MARGIN_LR, MARGIN_TOP)
            painter.setClipRect(0, 0, TEXT_WIDTH, TEXT_HEIGHT)
            doc.drawContents(painter, QtCore.QRectF(0, 0, TEXT_WIDTH, TEXT_HEIGHT))
            painter.restore()

            painter.end()

            if not _atomic_save_image(canvas, str(out_png)):
                raise RuntimeError("Failed to save message.png")

            self.status.setText("🔍 message.png generated.")
        except Exception as e:
            # Force existence (wall->message) if generation fails
            self.status.setText(f"❌ Error generating message.png: {e}")
            self._ensure_message_exists()
    def shutdown(self) -> None:
        self._overlay_render_timer.stop()
        try:
            if self._tab_active:
                self.deactivate_for_tab_change()
            else:
                self.sync_from_disk()
        except Exception:
            pass
